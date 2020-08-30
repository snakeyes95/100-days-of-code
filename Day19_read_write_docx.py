import docx

document=docx.Document('sample_text.docx')

all_paragraphs=document.paragraphs

for para in all_paragraphs:
    print(para.text)

single_para = document.paragraphs[0]
for run in single_para.runs:
    print(run.text)